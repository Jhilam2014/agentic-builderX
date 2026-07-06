from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

OUT = Path(__file__).resolve().parent
NAVY = "12233F"
BLUE = "2563EB"
TEAL = "0F766E"
INK = "172033"
MUTED = "64748B"
LIGHT = "F1F5F9"
PALE_BLUE = "EFF6FF"
PALE_GREEN = "ECFDF5"
PALE_GOLD = "FFF7E6"
RED = "BE123C"
WHITE = "FFFFFF"
DXA = 9360


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Aptos"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_document(doc, preset="business"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12 if preset == "business" else 1.2
    for style_name, size, before, after, color in [
        ("Title", 29, 0, 8, NAVY),
        ("Subtitle", 13, 0, 18, MUTED),
        ("Heading 1", 17, 17, 8, NAVY),
        ("Heading 2", 13, 12, 5, BLUE),
        ("Heading 3", 11, 9, 4, TEAL),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display" if style_name in ("Title", "Heading 1") else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = "Aptos"
        style.font.size = Pt(10.3)
        style.paragraph_format.left_indent = Inches(0.38)
        style.paragraph_format.first_line_indent = Inches(-0.19)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("AGENTIC BUILDERX")
    set_font(r, 8.5, BLUE, True)
    r = p.add_run("  /  BUSINESS DOCUMENT")
    set_font(r, 8.5, MUTED, False)
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Prepared 5 July 2026  |  Confidential budgetary material  |  ")
    set_font(r, 8, MUTED)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    p._p.append(fld)


def add_title_block(doc, kicker, title, subtitle, meta):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(kicker.upper())
    set_font(r, 9, TEAL, True)
    p = doc.add_paragraph(style="Title")
    p.add_run(title)
    p = doc.add_paragraph(style="Subtitle")
    p.add_run(subtitle)
    table = doc.add_table(rows=1, cols=len(meta))
    set_table_geometry(table, [DXA // len(meta)] * len(meta))
    for cell, (label, value) in zip(table.rows[0].cells, meta):
        set_cell_shading(cell, PALE_BLUE)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(label.upper() + "\n")
        set_font(r, 7.5, BLUE, True)
        r = p.add_run(value)
        set_font(r, 10, NAVY, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_callout(doc, label, text, fill=PALE_BLUE, accent=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    r = p.add_run(label.upper() + "  ")
    set_font(r, 8.5, accent, True)
    r = p.add_run(text)
    set_font(r, 10.5, INK, True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_table(doc, headers, rows, widths, header_fill=NAVY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    set_repeat_header(table.rows[0])
    for cell, header in zip(table.rows[0].cells, headers):
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p.add_run(header)
        set_font(r, 9, WHITE, True)
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for col_idx, (cell, value) in enumerate(zip(cells, values)):
            if row_idx % 2:
                set_cell_shading(cell, "F8FAFC")
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            set_font(r, 9.2, INK, col_idx == 0)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_source(doc, title, url):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f"{title}: {url}")
    set_font(r, 8.5, MUTED)


def page_break(doc):
    doc.add_page_break()


def build_quote():
    doc = Document()
    configure_document(doc, "business")
    add_title_block(
        doc,
        "Cloud hosting budget estimate",
        "Agentic BuilderX Production Platform",
        "A non-binding monthly operating estimate for an AWS Mumbai deployment.",
        [("Region", "AWS Mumbai"), ("Capacity", "25 concurrent builds"), ("Volume", "5,000 builds/month")],
    )
    add_callout(doc, "Expected monthly budget", "$3,400-$4,400 USD, including infrastructure and expected OpenAI model usage.", PALE_GREEN, TEAL)
    p = doc.add_paragraph()
    p.add_run("This estimate is designed for a multi-AZ launch serving approximately 100-300 active users. ")
    r = p.add_run("It is not a binding quotation or invoice.")
    r.bold = True
    doc.add_heading("Executive recommendation", level=1)
    doc.add_paragraph("Deploy BuilderX as a managed control plane with isolated, queue-driven build workers. The existing local Docker Compose topology is suitable for development, but it should not be exposed directly as a public multi-user service.")
    add_bullets(doc, [
        "Serve the BuilderX frontend and static previews through CloudFront and S3.",
        "Run the API as stateless ECS services across at least two availability zones.",
        "Queue builds in SQS and execute each build in an isolated, time-limited worker task.",
        "Store users, projects, builds, permissions, decisions, and snapshots in managed databases.",
        "Store generated files, media, exports, and workspace archives in S3.",
    ])
    doc.add_heading("Estimated monthly budget", level=1)
    add_table(doc, ["Cost category", "Expected range", "What it covers"], [
        ("API and control plane", "$90-$180", "Multi-AZ ECS services for the BuilderX API and orchestration control plane."),
        ("Build and preview workers", "$150-$400", "Approximately 5,000 isolated builds, queued overflow, and time-limited previews."),
        ("Aurora PostgreSQL", "$180-$350", "Project ownership, workflow state, build records, permissions, and snapshots."),
        ("Redis and queues", "$60-$140", "Live workflow state, event distribution, coordination, and SQS requests."),
        ("Network and load balancing", "$100-$250", "ALB, private connectivity, routing, and typical transfer."),
        ("Storage and delivery", "$50-$180", "S3, CloudFront, ECR, build artifacts, exports, and backups."),
        ("Security and operations", "$100-$300", "WAF, CloudWatch, audit logging, alarms, and backup operations."),
        ("Neo4j AuraDB", "About $292", "Business Critical minimum for the production agent relationship graph."),
        ("Infrastructure subtotal", "$970-$1,940", "Excludes model usage, engineering labour, taxes, and premium support."),
    ], [2600, 1700, 5060])
    page_break(doc)
    doc.add_heading("AI model usage", level=1)
    doc.add_paragraph("Model usage is the largest variable. The estimates below use published GPT-5 pricing of $1.25 per million input tokens and $10 per million output tokens. Actual coding-agent usage depends on task complexity, retries, reviewers, context size, and caching.")
    add_table(doc, ["Scenario", "Tokens per build", "Approx. monthly AI cost"], [
        ("Efficient", "75k input + 15k output", "$1,220"),
        ("Expected", "150k input + 30k output", "$2,438"),
        ("Heavy", "300k input + 60k output", "$4,875"),
    ], [2200, 3700, 3460], TEAL)
    add_callout(doc, "Cost-control opportunity", "Use smaller models for classification, routine edits, and low-risk review. Reserve GPT-5-class execution for complex or high-risk work. A hybrid route could reduce expected AI usage to roughly $700-$1,500 per month.", PALE_GOLD, TEAL)
    doc.add_heading("Required production work", level=1)
    add_numbered(doc, [
        "Remove Docker-socket access and replace host-created preview containers with isolated cloud tasks.",
        "Replace local JSON, JSONL, logs, workspaces, and project registries with managed persistence.",
        "Cryptographically verify Google identity tokens and enforce ownership on every project resource.",
        "Convert synchronous generation into durable queued jobs with status, cancellation, retries, and idempotency.",
        "Build production frontend assets; do not operate public Vite development servers.",
        "Add quotas, rate limits, audit logs, malware scanning, worker network controls, monitoring, and tested backups.",
    ])
    doc.add_heading("Architecture at a glance", level=1)
    add_table(doc, ["Layer", "Recommended service", "Responsibility"], [
        ("Edge", "Route 53, ACM, CloudFront, WAF", "DNS, TLS, caching, and perimeter protection."),
        ("Application", "ALB + ECS Fargate", "Authenticated API and orchestration control plane."),
        ("Build execution", "SQS + isolated ECS tasks", "Sandboxed agent execution with bounded resources."),
        ("Data", "Aurora PostgreSQL + Redis", "Durable metadata and short-lived coordination state."),
        ("Knowledge", "Neo4j AuraDB + vector store", "Agent relationships and reusable memory."),
        ("Artifacts", "S3 + CloudFront", "Generated files, exports, media, snapshots, and previews."),
    ], [1750, 2900, 4710])
    page_break(doc)
    doc.add_heading("Assumptions and exclusions", level=1)
    add_bullets(doc, [
        "AWS Mumbai, multi-AZ operation, 100-300 active users, and 25 concurrent builds.",
        "Approximately 5,000 build instructions per month and a ten-minute average worker duration.",
        "Static previews are preferred; dynamic previews expire after 30 minutes of inactivity.",
        "Amounts exclude implementation labour, GST or other taxes, premium support, unusually high egress, and third-party observability subscriptions.",
        "Provider rates and exchange rates may change; confirm the final design in the AWS Pricing Calculator before purchase.",
    ])
    add_callout(doc, "Validity", "Budgetary estimate valid for planning purposes for 30 days from 5 July 2026. Final costs are usage-based and charged directly by service providers.", LIGHT, MUTED)
    doc.add_heading("Reference pricing", level=1)
    add_source(doc, "Amazon Lightsail instance bundles", "https://docs.aws.amazon.com/lightsail/latest/userguide/amazon-lightsail-bundles.html")
    add_source(doc, "AWS Fargate pricing", "https://aws.amazon.com/fargate/pricing/")
    add_source(doc, "AWS Elastic Load Balancing pricing", "https://aws.amazon.com/elasticloadbalancing/pricing/")
    add_source(doc, "OpenAI GPT-5 developer pricing", "https://openai.com/index/introducing-gpt-5-for-developers/")
    add_source(doc, "Neo4j AuraDB pricing", "https://neo4j.com/pricing/")
    doc.add_heading("Budget decision", level=1)
    doc.add_paragraph("For a private beta, a single hardened virtual machine can reduce initial infrastructure cost significantly. For the selected 25-concurrent-build production target, the managed architecture above is the safer baseline because it separates users, build workers, data, and credentials.")
    path = OUT / "Agentic-BuilderX-Cloud-Hosting-Budget-Estimate.docx"
    doc.save(path)
    return path


def build_guide():
    doc = Document()
    configure_document(doc, "guide")
    add_title_block(
        doc,
        "Business guide",
        "How Agentic BuilderX Builds Applications",
        "From a user instruction to a validated application, with reusable agents, memory, and visible decisions.",
        [("Audience", "Business stakeholders"), ("Format", "10-stage workflow"), ("Outcome", "Traceable application build")],
    )
    add_callout(doc, "In one sentence", "BuilderX turns a business instruction into a governed build: it plans the work, assigns specialist agents, records selected and rejected decisions, validates evidence, and preserves the result for the next instruction.", PALE_GREEN, TEAL)
    doc.add_heading("The operating model", level=1)
    doc.add_paragraph("BuilderX is the parent orchestration authority. It owns the original objective, decides how much orchestration is necessary, delegates bounded work, and determines whether the completed result satisfies the request.")
    add_table(doc, ["Participant", "Purpose", "Decision authority"], [
        ("BuilderX Fullstack Agent", "Plans the workflow, controls scope, assigns work, validates evidence, and closes the build.", "Final authority"),
        ("Project orchestrator", "Applies project-specific policies, context, architecture, and memory.", "Bounded execution"),
        ("Execution agents", "Perform focused frontend, backend, data, runtime, content, or integration work.", "Assigned task only"),
        ("QAgents", "Detect gaps, challenge weak evidence, and propose the next instruction when needed.", "Review and correction"),
        ("Independent reviewer", "Checks workspace evidence without modifying the project.", "Pass/fail recommendation"),
        ("Human user", "Defines the objective and resolves choices that require business judgment.", "Business approval"),
    ], [2100, 4650, 2610])
    doc.add_heading("End-to-end build flow", level=1)
    add_numbered(doc, [
        "The user selects or creates a project and submits an instruction.",
        "BuilderX reads the task type, project identity, policies, available agents, and constraints.",
        "Adaptive orchestration selects a direct, delegated, or independently reviewed route.",
        "BuilderX creates a parent workflow and assigns bounded child executions to responsible agents.",
        "Agents implement selected features using project-specific context and reusable memory.",
        "Each decision point records the available options, selected branches, rejected branches, reasons, and responsible agent.",
        "BuilderX collects generated-file evidence, validation results, reviewer findings, and runtime readiness.",
        "Failures can be retried, simplified, redirected, or returned for human choice without falsely reporting completion.",
        "A successful build produces a preview, changed files, generated features, agent work records, and an execution snapshot.",
        "The project retains its history and memory so the next instruction begins with stronger context.",
    ])
    doc.add_heading("How adaptive orchestration chooses a route", level=1)
    doc.add_paragraph("Not every instruction needs the same number of agents or model calls. BuilderX evaluates complexity, risk, project ownership, validation requirements, and the available model-call budget.")
    add_table(doc, ["Route", "When it is selected", "Typical flow"], [
        ("Single", "Simple, low-risk work where delegation would add overhead.", "BuilderX plans, executes, validates, and records the result."),
        ("Delegated", "Project-specific work benefits from a local orchestrator or specialist executor.", "BuilderX delegates a bounded task and retains completion authority."),
        ("Delegated + reviewed", "Hard, high-risk, or validation-sensitive work needs independent evidence.", "Executor changes the workspace; a separate reviewer inspects it; BuilderX decides completion."),
    ], [1900, 3650, 3810], TEAL)
    add_callout(doc, "Why this matters", "Adaptive routing avoids paying for unnecessary agents on simple work while preserving stronger review for complex or risky builds.", PALE_BLUE, BLUE)
    doc.add_heading("What a decision tree records", level=1)
    doc.add_paragraph("Every build begins at one starting point and advances through staged decision points. Rejected choices terminate at the point where they were rejected. Selected choices continue into the next decision stage.")
    add_table(doc, ["Stage", "Illustrative selected branch", "Illustrative rejected branches"], [
        ("Route selection", "Delegated execution", "Single route; delegated + independent review"),
        ("Agent assignment", "Project orchestrator and frontend agent", "Unassigned execution"),
        ("Feature scope", "Media analysis, filtering, reports", "Template-only output; unrelated expansion"),
        ("Completion gate", "Approve build after evidence passes", "Reject completion when evidence is sufficient"),
    ], [1800, 3900, 3660])
    doc.add_heading("Selection and rejection evidence", level=2)
    add_bullets(doc, [
        "Every branch includes a state: selected, rejected, completed, passed, or failed.",
        "Every decision includes its reason, constraint, and responsible agent.",
        "Rejected branches remain visible so stakeholders understand what the system deliberately did not build.",
        "Selected feature branches identify what was generated and which execution agent produced the evidence.",
    ])
    doc.add_heading("What agents do during a build", level=1)
    doc.add_paragraph("Agents are reusable workers with distinct responsibilities and memory. Their identities remain visually consistent across the Agents page, runtime activity, decision tree, and Execution Snapshot.")
    add_table(doc, ["Work category", "Examples of recorded agent work"], [
        ("Planning", "Classify the task, preserve the objective, identify constraints, and choose an execution route."),
        ("Project context", "Load project policies, existing architecture, prior decisions, and relevant memory."),
        ("Implementation", "Create or modify routes, components, data structures, styles, services, tests, and configuration."),
        ("Feature generation", "Record each generated functionality and the action or file that provides evidence for it."),
        ("Validation", "Inspect files, tests, runtime state, preview readiness, and independent review results."),
        ("Learning", "Store durable lessons, correction patterns, and reusable project knowledge for future tasks."),
    ], [2400, 6960])
    doc.add_heading("Project memory and reusable expertise", level=1)
    add_bullets(doc, [
        "Global agent knowledge provides reusable domain capability across projects.",
        "Project memory preserves local policies, architecture, prior builds, decisions, and correction history.",
        "Vector memory supports semantic retrieval of relevant knowledge rather than loading everything into every prompt.",
        "Graph memory describes relationships between agents, projects, capabilities, decisions, and outputs.",
        "Memory remains advisory: BuilderX still preserves the current user instruction as the parent objective.",
    ])
    add_callout(doc, "Efficiency principle", "Reuse the smallest relevant context and the fewest agents needed for the task. Add delegation or review only when it materially improves accuracy, risk control, or validation.", PALE_GOLD, TEAL)
    doc.add_heading("The Execution Snapshot", level=1)
    doc.add_paragraph("Each successful or failed build receives an immutable execution snapshot. The snapshot is a historical record, not merely a live animation.")
    add_bullets(doc, [
        "Build and parent workflow identity, start time, completion time, and duration.",
        "Chronological execution events and child-agent assignments.",
        "Selected and rejected decisions with reasons and constraints.",
        "Responsible agent for every decision and generated feature.",
        "Agent work summaries, changed files, validation state, review result, and failure evidence.",
    ])
    page_break(doc)
    doc.add_heading("Success, failure, and retry", level=1)
    add_table(doc, ["Outcome", "What BuilderX does", "What the user sees"], [
        ("Successful", "Confirms execution and validation evidence, records completion, and publishes the preview.", "Successful build, features, files, agent work, and full snapshot."),
        ("Execution failure", "Records partial evidence, marks completion as rejected, and applies retry policy only when safe.", "Failed snapshot with cause, completed work, and recovery choices."),
        ("Validation failure", "Prevents false completion even when files were generated.", "Rejected completion and the evidence that did not pass."),
        ("Human decision needed", "Pauses automated progression and presents bounded recovery choices.", "Retry, simplify scope, or change architecture."),
    ], [1700, 4300, 3360], NAVY)
    doc.add_heading("What the user receives", level=1)
    add_bullets(doc, [
        "A project-specific application preview.",
        "A durable build ID and parent workflow ID.",
        "Generated and modified files.",
        "A list of selected features and deliberately rejected alternatives.",
        "Agent identities and concise work performed toward the build.",
        "A movable decision graph with a persistent evidence panel.",
        "A project history that can be revisited build by build.",
    ])
    doc.add_heading("Production hosting model", level=1)
    doc.add_paragraph("For cloud operation, the BuilderX control plane should be separated from isolated build workers. User instructions enter a durable queue; workers receive temporary workspaces and bounded credentials; artifacts and snapshots are saved before workers terminate.")
    add_table(doc, ["Boundary", "Production behavior"], [
        ("Identity", "Verified Google identity or enterprise authentication with server-enforced ownership."),
        ("Control plane", "Stateless API and orchestration services running across multiple availability zones."),
        ("Build workers", "Isolated, non-privileged tasks with CPU, memory, time, and network limits."),
        ("Persistence", "Managed databases and object storage instead of local JSON files and host directories."),
        ("Previews", "Static publication by default; authenticated, expiring dynamic previews when required."),
        ("Secrets", "Short-lived access through a managed secrets service; no developer credentials mounted into workers."),
    ], [2100, 7260], TEAL)
    page_break(doc)
    doc.add_heading("Business value", level=1)
    add_table(doc, ["Capability", "Business benefit"], [
        ("Visible decisions", "Stakeholders can see not only what was built, but what was rejected and why."),
        ("Reusable agents", "Specialist capability improves over time instead of being recreated for every project."),
        ("Project memory", "Future work starts from established context, reducing repeated discovery and inconsistent decisions."),
        ("Adaptive cost", "Simple tasks stay lightweight while difficult tasks receive stronger delegation and review."),
        ("Evidence-based completion", "Generated files alone do not count as success; validation evidence controls approval."),
        ("Build snapshots", "Successes and failures remain explainable, auditable, and comparable over time."),
    ], [2300, 7060])
    add_callout(doc, "Core promise", "Agentic BuilderX is designed to make application-building decisions traceable, reusable, and governable without removing the user from important business choices.", PALE_GREEN, TEAL)
    doc.add_heading("Glossary", level=1)
    add_table(doc, ["Term", "Meaning"], [
        ("Parent workflow", "The BuilderX-owned record representing one complete user instruction."),
        ("Child execution", "A bounded task assigned to a project or specialist agent."),
        ("Adaptive route", "The selected orchestration depth: single, delegated, or independently reviewed."),
        ("QAgent", "A quality-focused agent that detects objective gaps and proposes corrective next work."),
        ("Execution Snapshot", "The immutable record of events, decisions, agents, outputs, and validation for one build."),
        ("Completion gate", "The final BuilderX decision that approves or rejects workflow completion."),
    ], [2300, 7060], NAVY)
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("A BuilderX build is a governed sequence rather than a single model response. The platform preserves the user objective, assigns accountable agents, makes branching decisions visible, validates evidence, and carries project knowledge forward into the next instruction.")
    path = OUT / "How-Agentic-BuilderX-Builds-Applications.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    print(build_quote())
    print(build_guide())
